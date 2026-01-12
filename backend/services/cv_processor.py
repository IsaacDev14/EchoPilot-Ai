"""
EchoPilot CV/Resume Processing Service
Handles extraction of text from PDF, DOCX, and TXT files.
"""

import io
from typing import Optional, Tuple
from PyPDF2 import PdfReader
from docx import Document


class CVProcessor:
    """Process and extract text from CV/Resume files."""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}
    
    @staticmethod
    async def extract_text(file_content: bytes, filename: str) -> Tuple[str, Optional[str]]:
        """
        Extract text from a file.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename to determine type
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        extension = "." + filename.lower().split(".")[-1] if "." in filename else ""
        
        if extension not in CVProcessor.SUPPORTED_EXTENSIONS:
            return "", f"Unsupported file type: {extension}"
        
        try:
            if extension == ".pdf":
                return CVProcessor._extract_from_pdf(file_content), None
            elif extension in {".docx", ".doc"}:
                return CVProcessor._extract_from_docx(file_content), None
            elif extension == ".txt":
                return file_content.decode("utf-8", errors="ignore"), None
            else:
                return "", f"Unsupported file type: {extension}"
        except Exception as e:
            return "", f"Error extracting text: {str(e)}"
    
    @staticmethod
    def _extract_from_pdf(content: bytes) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def _extract_from_docx(content: bytes) -> str:
        """Extract text from DOCX file."""
        doc = Document(io.BytesIO(content))
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    
    @staticmethod
    def create_summary(text: str, max_length: int = 2000) -> str:
        """
        Create a summarized version of the CV text for context.
        This is a simple truncation - in production, you might use an LLM.
        """
        if len(text) <= max_length:
            return text
        
        # Try to truncate at a sentence boundary
        truncated = text[:max_length]
        last_period = truncated.rfind(".")
        if last_period > max_length * 0.7:
            return truncated[:last_period + 1]
        
        return truncated + "..."


# Global CV context storage (in production, use Redis or database)
class CVContextManager:
    """Manages CV context for sessions."""
    
    def __init__(self):
        self._contexts = {}
    
    def set_context(self, session_id: str, filename: str, text: str, summary: str = None):
        """Store CV context for a session."""
        from datetime import datetime
        self._contexts[session_id] = {
            "filename": filename,
            "text": text,
            "summary": summary or CVProcessor.create_summary(text),
            "uploaded_at": datetime.utcnow()
        }
    
    def get_context(self, session_id: str) -> Optional[dict]:
        """Get CV context for a session."""
        return self._contexts.get(session_id)
    
    def clear_context(self, session_id: str):
        """Clear CV context for a session."""
        if session_id in self._contexts:
            del self._contexts[session_id]
    
    def has_context(self, session_id: str) -> bool:
        """Check if session has CV context."""
        return session_id in self._contexts


# Global instance
cv_context_manager = CVContextManager()
